import docx
import openpyxl
import pandas as pd
from typing import List, Dict
import logging
import os

logger = logging.getLogger(__name__)


class DocumentParser:
    """Base class for document parsing."""

    @staticmethod
    def parse_docx(file_path: str) -> Dict:
        """
        Parse DOCX file and extract text content.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dict with document metadata and text content
        """
        try:
            doc = docx.Document(file_path)
            text_blocks = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_blocks.append(para.text.strip())

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_blocks.append(row_text)

            full_text = '\n\n'.join(text_blocks)

            return {
                'document_name': os.path.basename(file_path),
                'document_type': 'docx',
                'full_text': full_text,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }

        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {e}")
            raise

    @staticmethod
    def parse_xlsx(file_path: str) -> Dict:
        """
        Parse XLSX file and extract structured data.

        Args:
            file_path: Path to the XLSX file

        Returns:
            Dict with document metadata and structured content
        """
        try:
            # Read Excel file
            df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets

            text_blocks = []
            all_data = []

            for sheet_name, sheet_df in df.items():
                # Add sheet name as context
                text_blocks.append(f"--- {sheet_name} ---")

                # Convert each row to text
                for idx, row in sheet_df.iterrows():
                    # Create readable text from row
                    row_text = ' | '.join([
                        f"{col}: {val}"
                        for col, val in row.items()
                        if pd.notna(val) and str(val).strip()
                    ])

                    if row_text:
                        text_blocks.append(row_text)
                        all_data.append(row.to_dict())

            full_text = '\n'.join(text_blocks)

            return {
                'document_name': os.path.basename(file_path),
                'document_type': 'xlsx',
                'full_text': full_text,
                'sheet_count': len(df),
                'row_count': sum(len(sheet_df) for sheet_df in df.values()),
                'structured_data': all_data
            }

        except Exception as e:
            logger.error(f"Error parsing XLSX file {file_path}: {e}")
            raise

    @staticmethod
    def extract_metadata_from_xlsx(file_path: str) -> Dict:
        """
        Extract specific metadata from process XLSX files.

        Returns department, process owner, etc.
        """
        try:
            df = pd.read_excel(file_path, sheet_name=0)

            # Extract department from filename
            filename = os.path.basename(file_path)
            department = filename.replace('Zhodnocení procesů ', '').replace('.xlsx', '')

            metadata = {
                'department': department,
                'process_owner': None,
                'process_count': len(df)
            }

            # Try to extract process owner if column exists
            if 'Vlastník procesu' in df.columns:
                owners = df['Vlastník procesu'].dropna().unique()
                metadata['process_owner'] = ', '.join(owners[:3]) if len(owners) > 0 else None

            return metadata

        except Exception as e:
            logger.warning(f"Error extracting metadata from {file_path}: {e}")
            return {'department': None, 'process_owner': None}


def parse_document(file_path: str) -> Dict:
    """
    Parse any supported document type.

    Args:
        file_path: Path to the document

    Returns:
        Parsed document data
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.docx' or ext == '.doc':
        return DocumentParser.parse_docx(file_path)
    elif ext == '.xlsx':
        parsed = DocumentParser.parse_xlsx(file_path)
        metadata = DocumentParser.extract_metadata_from_xlsx(file_path)
        parsed['metadata'] = metadata
        return parsed
    else:
        raise ValueError(f"Unsupported file type: {ext}")
